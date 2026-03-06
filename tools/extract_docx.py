from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


def _norm_ws(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def extract_docx_to_text(docx_path: Path) -> str:
    d = Document(str(docx_path))
    out: list[str] = []

    # Paragraphs
    for p in d.paragraphs:
        t = _norm_ws(p.text)
        if t:
            out.append(t)

    # Tables
    if d.tables:
        out.append("")
    for ti, table in enumerate(d.tables, start=1):
        out.append(f"[TABLE {ti}]")
        for row in table.rows:
            cells = [_norm_ws(c.text) for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
        out.append("")

    return "\n".join(out).rstrip() + "\n"


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: extract_docx.py <input.docx> [output.txt]", file=sys.stderr)
        return 2

    in_path = Path(argv[1]).expanduser().resolve()
    if not in_path.exists():
        print(f"Input file not found: {in_path}", file=sys.stderr)
        return 2

    text = extract_docx_to_text(in_path)

    if len(argv) >= 3:
        out_path = Path(argv[2]).expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
    else:
        sys.stdout.write(text)

    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

